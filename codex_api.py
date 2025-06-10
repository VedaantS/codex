from flask import Flask, request, jsonify, abort, send_from_directory, render_template
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from flask_jwt_extended import (
    JWTManager, create_access_token, jwt_required, get_jwt_identity
)
from werkzeug.security import generate_password_hash, check_password_hash
from datetime import timedelta
import boto3
import os
from flask_cors import CORS
import uuid
import shutil
import datetime
import re
import openai
from PyPDF2 import PdfReader
from docx import Document
import difflib

app = Flask(__name__, static_folder='.', static_url_path='')
CORS(app)

# Configuration
database_url = os.getenv('DATABASE_URL', 'postgresql://vedaantsrivastava:mypassword@localhost/vedaantsrivastava')
app.config['SQLALCHEMY_DATABASE_URI'] = database_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['JWT_SECRET_KEY'] = 'c2uWLdF3Do_24HeDnAzlv7zkrrCfJJU69igcm_fiiKU'
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = timedelta(hours=1)
app.config['S3_BUCKET'] = os.getenv('S3_BUCKET')

UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Extensions
db = SQLAlchemy(app)
migrate = Migrate(app, db)
jwt = JWTManager(app)
s3 = boto3.client('s3')

# -- Models --
class User(db.Model):
    __tablename__ = 'users'
    id = db.Column(db.String, primary_key=True)
    email = db.Column(db.String(255), unique=True, nullable=False)
    name = db.Column(db.String(255), nullable=False)
    password_hash = db.Column(db.String(255), nullable=False)
    role = db.Column(db.Enum('scientist', 'funder', 'admin', name='role_enum'), nullable=False)
    avatar_url = db.Column(db.Text)
    created_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    updated_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now(), onupdate=db.func.now())
    profile = db.relationship('Profile', uselist=False, back_populates='user')

    def set_password(self, password):
        # Use pbkdf2:sha256 explicitly to avoid scrypt issues
        self.password_hash = generate_password_hash(password, method='pbkdf2:sha256')

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class Profile(db.Model):
    __tablename__ = 'profiles'
    user_id = db.Column(db.String, db.ForeignKey('users.id'), primary_key=True)
    bio = db.Column(db.Text)
    affiliation = db.Column(db.String(255))
    expertise_tags = db.Column(db.ARRAY(db.String))
    user = db.relationship('User', back_populates='profile')

class Experiment(db.Model):
    __tablename__ = 'experiments'
    id = db.Column(db.String, primary_key=True)
    title = db.Column(db.String(255), nullable=False)
    description = db.Column(db.Text)
    owner_id = db.Column(db.String, db.ForeignKey('users.id'), nullable=False)
    visibility = db.Column(db.Enum('public', 'collaborators_only', 'private', name='visibility_enum'), nullable=False, server_default='private')
    created_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    updated_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now(), onupdate=db.func.now())
    owner = db.relationship('User', backref='experiments')

class ProtocolVersion(db.Model):
    __tablename__ = 'protocol_versions'
    id = db.Column(db.String, primary_key=True)
    experiment_id = db.Column(db.String, db.ForeignKey('experiments.id'), nullable=False)
    version_label = db.Column(db.String(50), nullable=False)
    parent_version_id = db.Column(db.String, db.ForeignKey('protocol_versions.id'), nullable=True)
    metadata_ = db.Column(db.JSON)
    created_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    step_map = db.Column(db.JSON)  # <-- New field for flowchart persistence
    experiment = db.relationship('Experiment', backref='protocol_versions')
    parent = db.relationship('ProtocolVersion', remote_side=[id])

class ExperimentStep(db.Model):
    __tablename__ = 'experiment_steps'
    id = db.Column(db.String, primary_key=True)
    protocol_version_id = db.Column(db.String, db.ForeignKey('protocol_versions.id'), nullable=False)
    title = db.Column(db.String(255), nullable=False)
    due_date = db.Column(db.Date)
    done = db.Column(db.Boolean, default=False)
    content_markdown = db.Column(db.Text)
    results_markdown = db.Column(db.Text)  # <-- Add this line
    estimated_time_minutes = db.Column(db.Integer)
    assigned_to_id = db.Column(db.String, db.ForeignKey('users.id'))
    reproducibility_score = db.Column(db.Float)
    impact_score = db.Column(db.Float)
    difficulty_score = db.Column(db.Float)
    order_index = db.Column(db.Integer, nullable=False)
    created_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    updated_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now(), onupdate=db.func.now())
    protocol_version = db.relationship('ProtocolVersion', backref='steps')
    assigned_to = db.relationship('User')

    def to_dict(self):
        return {
            'id': self.id,
            'protocol_version_id': self.protocol_version_id,
            'title': self.title,
            'due_date': self.due_date.isoformat() if self.due_date else None,
            'done': self.done,
            'content_markdown': self.content_markdown,
            'results_markdown': self.results_markdown,  # <-- Add this line
            'estimated_time_minutes': self.estimated_time_minutes,
            'assigned_to_id': self.assigned_to_id,
            'reproducibility_score': self.reproducibility_score,
            'impact_score': self.impact_score,
            'difficulty_score': self.difficulty_score,
            'order_index': self.order_index,
            'created_at': self.created_at.isoformat() if self.created_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

class FileAttachment(db.Model):
    __tablename__ = 'file_attachments'
    id = db.Column(db.String, primary_key=True)
    owner_id = db.Column(db.String, db.ForeignKey('users.id'), nullable=False)
    experiment_step_id = db.Column(db.String, db.ForeignKey('experiment_steps.id'), nullable=False)
    filename = db.Column(db.String(255), nullable=False)
    storage_path = db.Column(db.Text, nullable=False)
    mime_type = db.Column(db.String(100))
    size_bytes = db.Column(db.BigInteger)
    uploaded_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    owner = db.relationship('User')
    step = db.relationship('ExperimentStep')

class ChatChannel(db.Model):
    __tablename__ = 'chat_channels'
    id = db.Column(db.String, primary_key=True)
    experiment_id = db.Column(db.String, db.ForeignKey('experiments.id'), nullable=False)
    name = db.Column(db.String(100), nullable=False)
    created_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    experiment = db.relationship('Experiment')

class ChatMessage(db.Model):
    __tablename__ = 'chat_messages'
    id = db.Column(db.String, primary_key=True)
    channel_id = db.Column(db.String, db.ForeignKey('chat_channels.id'), nullable=False)
    sender_id = db.Column(db.String, db.ForeignKey('users.id'), nullable=False)
    content = db.Column(db.Text, nullable=False)
    sent_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    channel = db.relationship('ChatChannel', backref='messages')
    sender = db.relationship('User')

class Grant(db.Model):
    __tablename__ = 'grants'
    id = db.Column(db.String, primary_key=True)
    title = db.Column(db.String(255), nullable=False)
    description = db.Column(db.Text)
    total_funding_usd = db.Column(db.Numeric(15,2))
    application_questions = db.Column(db.ARRAY(db.Text))
    created_by_id = db.Column(db.String, db.ForeignKey('users.id'), nullable=False)
    created_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    creator = db.relationship('User')

class Project(db.Model):
    __tablename__ = 'projects'
    id = db.Column(db.String, primary_key=True)
    owner_id = db.Column(db.String, db.ForeignKey('users.id'), nullable=False)
    title = db.Column(db.String(255), nullable=False)
    budget_requested = db.Column(db.Numeric(15,2))
    reproducibility_score = db.Column(db.Float)
    impact_score = db.Column(db.Float)
    difficulty_score = db.Column(db.Float)
    created_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    updated_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now(), onupdate=db.func.now())
    owner = db.relationship('User')
    paper_content = db.Column(db.Text)

class GrantApplication(db.Model):
    __tablename__ = 'grant_applications'
    id = db.Column(db.String, primary_key=True)
    grant_id = db.Column(db.String, db.ForeignKey('grants.id'), nullable=False)
    project_id = db.Column(db.String, db.ForeignKey('projects.id'), nullable=False)
    answers = db.Column(db.JSON)
    status = db.Column(db.Enum('pending','shortlisted','awarded','rejected', name='app_status_enum'), nullable=False, server_default='pending')
    submitted_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    updated_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now(), onupdate=db.func.now())
    grant = db.relationship('Grant')
    project = db.relationship('Project')

    def to_dict(self):
        return {
            'id': self.id,
            'grant_id': self.grant_id,
            'project_id': self.project_id,
            'answers': self.answers,
            'status': self.status,
            'submitted_at': self.submitted_at.isoformat() if self.submitted_at else None,
            'updated_at': self.updated_at.isoformat() if self.updated_at else None
        }

class Award(db.Model):
    __tablename__ = 'awards'
    id = db.Column(db.String, primary_key=True)
    application_id = db.Column(db.String, db.ForeignKey('grant_applications.id'), nullable=False)
    awarded_amount = db.Column(db.Numeric(15,2), nullable=False)
    awarded_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    application = db.relationship('GrantApplication')

class GrantMilestone(db.Model):
    __tablename__ = 'grant_milestones'
    id = db.Column(db.String, primary_key=True)
    award_id = db.Column(db.String, db.ForeignKey('awards.id'), nullable=False)
    name = db.Column(db.String(255), nullable=False)
    due_date = db.Column(db.Date)
    completed = db.Column(db.Boolean, nullable=False, server_default='false')
    completed_at = db.Column(db.DateTime(timezone=True))
    award = db.relationship('Award')

class DiscoveryItem(db.Model):
    __tablename__ = 'discovery_items'
    id = db.Column(db.String, primary_key=True)
    title = db.Column(db.String(255), nullable=False)
    description = db.Column(db.Text)
    field = db.Column(db.String(100))
    status = db.Column(db.String(100))
    lead_name = db.Column(db.String(255))
    tags = db.Column(db.ARRAY(db.String))
    ai_score = db.Column(db.Float)
    created_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())

class CollaborationSuggestion(db.Model):
    __tablename__ = 'collaboration_suggestions'
    id = db.Column(db.String, primary_key=True)
    for_user_id = db.Column(db.String, db.ForeignKey('users.id'), nullable=False)
    suggested_user_id = db.Column(db.String, db.ForeignKey('users.id'), nullable=False)
    reason = db.Column(db.Text)
    created_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    for_user = db.relationship('User', foreign_keys=[for_user_id])
    suggested_user = db.relationship('User', foreign_keys=[suggested_user_id])

class GlobalChatMessage(db.Model):
    __tablename__ = 'global_chat_messages'
    id = db.Column(db.String, primary_key=True)
    sender_id = db.Column(db.String, db.ForeignKey('users.id'), nullable=False)
    recipient_id = db.Column(db.String, db.ForeignKey('users.id'), nullable=False)
    content = db.Column(db.Text, nullable=False)
    sent_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    sender = db.relationship('User', foreign_keys=[sender_id])
    recipient = db.relationship('User', foreign_keys=[recipient_id])

class NotebookEntry(db.Model):
    __tablename__ = 'notebook_entries'
    id = db.Column(db.String, primary_key=True)
    project_id = db.Column(db.String, db.ForeignKey('projects.id'), nullable=False)
    user_id = db.Column(db.String, db.ForeignKey('users.id'))
    user_name = db.Column(db.String(255))
    timestamp = db.Column(db.DateTime(timezone=True), default=db.func.now())
    device = db.Column(db.String(255))
    location = db.Column(db.String(255))
    session_id = db.Column(db.String(255))
    experiment_id = db.Column(db.String(255))
    version = db.Column(db.String(50))
    visibility = db.Column(db.String(50), default='team')
    content = db.Column(db.Text)
    structured = db.Column(db.JSON)
    diffs = db.Column(db.JSON)
    attachments = db.relationship('NotebookAttachment', backref='entry', lazy=True)

    def to_dict(self):
        return {
            'id': self.id,
            'project_id': self.project_id,
            'user_id': self.user_id,
            'user_name': self.user_name,
            'timestamp': self.timestamp.isoformat() if self.timestamp else None,
            'device': self.device,
            'location': self.location,
            'session_id': self.session_id,
            'experiment_id': self.experiment_id,
            'version': self.version,
            'visibility': self.visibility,
            'content': self.content,
            'structured': self.structured,
            'diffs': self.diffs,
            'attachments': [a.to_dict() for a in self.attachments]
        }

class NotebookAttachment(db.Model):
    __tablename__ = 'notebook_attachments'
    id = db.Column(db.String, primary_key=True)
    entry_id = db.Column(db.String, db.ForeignKey('notebook_entries.id'))
    filename = db.Column(db.String(255))
    storage_path = db.Column(db.String(255))
    uploaded_at = db.Column(db.DateTime(timezone=True), default=db.func.now())
    def to_dict(self):
        return {
            'id': self.id,
            'filename': self.filename,
            'uploaded_at': self.uploaded_at.isoformat() if self.uploaded_at else None
        }

# -- Utility --
def get_model_or_404(model, id):
    instance = model.query.get(id)
    if not instance:
        abort(404, description=f'{model.__name__} not found')
    return instance

def experiment_to_dict(e):
    return {
        'id': e.id,
        'title': e.title,
        'description': e.description,
        'owner_id': e.owner_id,
        'visibility': e.visibility,
        'created_at': e.created_at.isoformat() if e.created_at else None,
        'updated_at': e.updated_at.isoformat() if e.updated_at else None,
    }

def protocol_version_to_dict(v):
    return {
        'id': v.id,
        'experiment_id': v.experiment_id,
        'version_label': v.version_label,
        'parent_version_id': v.parent_version_id,
        'metadata_': v.metadata_,
        'created_at': v.created_at.isoformat() if v.created_at else None,
    }


def proj_to_experiment(eid):
    proj = Project.query.get(eid)
    if proj:
        exp = Experiment.query.filter_by(title=proj.title, owner_id=proj.owner_id).first()
    return exp.id if exp else None

# -- Auth Routes --

@app.route('/auth/register', methods=['POST'])
def register():
    data = request.get_json()
    if User.query.filter_by(email=data['email']).first(): abort(400, 'Email already registered')
    user = User(id=db.func.gen_random_uuid(), email=data['email'], name=data['name'], role=data['role'])
    user.set_password(data['password'])
    db.session.add(user); db.session.commit()
    return jsonify(user_id=user.id), 201

@app.route('/auth/login', methods=['POST'])
def login():
    data = request.get_json(); user = User.query.filter_by(email=data['email']).first()
    if not user or not user.check_password(data['password']): abort(401, 'Invalid credentials')
    token = create_access_token(identity=user.id)
    return jsonify(access_token=token)

# -- User & Profile --
@app.route('/users/<id>', methods=['GET'])
@jwt_required()
def read_user(id):
    user = get_model_or_404(User, id)
    # Only return serializable fields
    return jsonify({
        'id': user.id,
        'email': user.email,
        'name': user.name,
        'role': user.role,
        'avatar_url': user.avatar_url,
        'created_at': user.created_at.isoformat() if user.created_at else None,
        'updated_at': user.updated_at.isoformat() if user.updated_at else None
    })

@app.route('/profiles/<id>', methods=['GET','PUT'])
def manage_profile(id):
    user = User.query.get(id)
    if not user:
        abort(404, description='User not found')
    if request.method=='GET':
        if user.profile:
            return jsonify({
                'user_id': user.profile.user_id,
                'bio': user.profile.bio,
                'affiliation': user.profile.affiliation,
                'expertise_tags': user.profile.expertise_tags or [],
            })
        else:
            return jsonify({'user_id': id, 'bio': '', 'affiliation': '', 'expertise_tags': []})
    # Only allow PUT if authenticated
    from flask_jwt_extended import verify_jwt_in_request, get_jwt_identity
    verify_jwt_in_request()
    data = request.get_json()
    if not user.profile:
        user.profile = Profile(user_id=id)
        db.session.add(user.profile)
    user.profile.bio = data.get('bio', user.profile.bio)
    user.profile.affiliation = data.get('affiliation', user.profile.affiliation)
    user.profile.expertise_tags = data.get('expertise_tags', user.profile.expertise_tags)
    db.session.commit()
    return jsonify(msg='updated')

# -- Users List & Query Endpoint --
@app.route('/users', methods=['GET'])
@jwt_required()
def users():
    email = request.args.get('email')
    if email:
        user = User.query.filter_by(email=email).first()
        if not user:
            return jsonify([]), 200
        return jsonify([{
            'id': user.id,
            'email': user.email,
            'name': user.name,
            'role': user.role,
            'avatar_url': user.avatar_url,
            'created_at': user.created_at.isoformat() if user.created_at else None,
            'updated_at': user.updated_at.isoformat() if user.updated_at else None
        }])
    # List all users (admin only in real app, but open for now)
    users = User.query.all()
    return jsonify([
        {
            'id': u.id,
            'email': u.email,
            'name': u.name,
            'role': u.role,
            'avatar_url': u.avatar_url,
            'created_at': u.created_at.isoformat() if u.created_at else None,
            'updated_at': u.updated_at.isoformat() if u.updated_at else None
        } for u in users
    ])

# -- Experiments CRUD --
@app.route('/experiments', methods=['GET','POST'])
def experiments():
    if request.method=='GET':
        user_id = request.args.get('user_id')
        if user_id:
            return jsonify([experiment_to_dict(e) for e in Experiment.query.filter_by(owner_id=user_id).all()])
        return jsonify([experiment_to_dict(e) for e in Experiment.query.all()])
    # Only allow POST if authenticated
    from flask_jwt_extended import verify_jwt_in_request, get_jwt_identity
    verify_jwt_in_request()
    data=request.get_json(); exp=Experiment(id=str(uuid.uuid4()), **data)
    db.session.add(exp)
    db.session.commit(); return jsonify(id=exp.id),201

@app.route('/experiments/<id>', methods=['GET','PATCH','DELETE'])
@jwt_required()
def experiment_detail(id):
    exp=get_model_or_404(Experiment,id)
    if request.method=='GET': 
        return jsonify(experiment_to_dict(exp))
    if request.method=='PATCH':
        for k,v in request.get_json().items(): setattr(exp,k,v)
        db.session.commit(); return jsonify(msg='updated')
    db.session.delete(exp); db.session.commit(); return '',204

# -- Protocol Versions & Steps --
@app.route('/experiments/<eid>/versions', methods=['GET','POST'])
@jwt_required()
def versions(eid):
    if request.method=='GET':
        return jsonify([protocol_version_to_dict(v) for v in ProtocolVersion.query.filter_by(experiment_id=eid)])
    data=request.get_json(); v=ProtocolVersion(id=db.func.gen_random_uuid(), experiment_id=eid, **data)
    db.session.add(v); db.session.commit(); return jsonify(id=v.id),201

@app.route('/versions/<vid>/steps', methods=['GET','POST'])
@jwt_required()
def steps(vid):
    if request.method=='GET':
        steps = ExperimentStep.query.filter_by(protocol_version_id=vid).all()
        import sys
        print(f"[DEBUG] /versions/{vid}/steps found {len(steps)} steps for protocol_version_id={vid} at {datetime.datetime.now()}", file=sys.stderr, flush=True)
        return jsonify([s.to_dict() for s in steps])
    data=request.get_json(); s=ExperimentStep(id=db.func.gen_random_uuid(), protocol_version_id=vid, **data)
    db.session.add(s); db.session.commit(); return jsonify(id=s.id),201

@app.route('/steps/<id>', methods=['GET','PATCH','DELETE'])
@jwt_required()
def step_detail(id):
    s = get_model_or_404(ExperimentStep, id)
    if request.method == 'GET':
        return jsonify(s.to_dict())
    if request.method == 'PATCH':
        data = request.get_json()
        # Only update fields present in the request, preserve others
        for k, v in data.items():
            if hasattr(s, k) and v is not None:
                print("YOOOOOO", s, k, v)
                setattr(s, k, v)
        db.session.commit()
        return jsonify(msg='updated')
    db.session.delete(s)
    db.session.commit()
    return '', 204

# -- File Upload URL --
@app.route('/steps/<id>/files/upload-url', methods=['POST'])
@jwt_required()
def upload_url(id):
    filename=request.json['filename']; key=f"{id}/{filename}"
    url=s3.generate_presigned_url('put_object', Params={'Bucket':app.config['S3_BUCKET'],'Key':key}, ExpiresIn=3600)
    return jsonify(upload_url=url, storage_path=key)

@app.route('/steps/<id>/attachments', methods=['POST'])
@jwt_required()
def upload_step_attachment(id):
    s = get_model_or_404(ExperimentStep, id)
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    filename = file.filename
    ext = os.path.splitext(filename)[1]
    unique_id = str(uuid.uuid4())
    storage_filename = f"{unique_id}{ext}"
    storage_path = os.path.join(UPLOAD_FOLDER, storage_filename)
    file.save(storage_path)
    fa = FileAttachment(
        id=unique_id,
        owner_id=get_jwt_identity(),
        experiment_step_id=id,
        filename=filename,
        storage_path=storage_filename,
        mime_type=file.mimetype,
        size_bytes=os.path.getsize(storage_path)
    )
    db.session.add(fa)
    db.session.commit()
    return jsonify({'id': fa.id, 'filename': fa.filename}), 201

@app.route('/steps/<id>/attachments', methods=['GET'])
@jwt_required()
def list_step_attachments(id):
    attachments = FileAttachment.query.filter_by(experiment_step_id=id).all()
    return jsonify([
        {
            'id': a.id,
            'filename': a.filename,
            'uploaded_at': a.uploaded_at.isoformat() if a.uploaded_at else None,
            'size_bytes': a.size_bytes,
            'mime_type': a.mime_type
        } for a in attachments
    ])

@app.route('/attachments/<attachment_id>/download', methods=['GET'])
def download_attachment(attachment_id):
    a = get_model_or_404(FileAttachment, attachment_id)
    return send_from_directory(UPLOAD_FOLDER, a.storage_path, as_attachment=True, download_name=a.filename)

# -- Chat --
@app.route('/experiments/<eid>/chat', methods=['GET','POST'])
@jwt_required()
def chat(eid):
    # Try direct experiment lookup
    exp = Experiment.query.get(eid)
    if not exp:
        # Try mapping project ID to experiment by matching title and owner_id
        proj = Project.query.get(eid)
        if proj:
            exp = Experiment.query.filter_by(title=proj.title, owner_id=proj.owner_id).first()
    if not exp:
        abort(404, description='Experiment not found')
    ch = ChatChannel.query.filter_by(experiment_id=exp.id).first()
    if not ch:
        ch = ChatChannel(id=db.func.gen_random_uuid(), experiment_id=exp.id, name='default')
        db.session.add(ch)
        db.session.commit()
    if request.method=='GET':
        def serialize_msg(m):
            return {
                'id': m.id,
                'channel_id': m.channel_id,
                'sender_id': m.sender_id,
                'content': m.content,
                'sent_at': m.sent_at.isoformat() if m.sent_at else None,
                'sender_name': User.query.get(m.sender_id).name if m.sender_id else 'Unknown',
            }
        msgs = [serialize_msg(m) for m in ch.messages]
        return jsonify(msgs)
    data = request.get_json()
    msg = ChatMessage(id=db.func.gen_random_uuid(), channel_id=ch.id, sender_id=get_jwt_identity(), content=data['content'])
    db.session.add(msg)
    db.session.commit()
    return jsonify(id=msg.id),201

@app.route('/global-chat/<user1_id>/<user2_id>', methods=['GET', 'POST'])
@jwt_required()
def global_chat(user1_id, user2_id):
    # Only allow access if one of the users is the current user
    current_user = get_jwt_identity()
    if current_user not in [user1_id, user2_id]:
        return jsonify({'error': 'Unauthorized'}), 403
    if request.method == 'GET':
        msgs = GlobalChatMessage.query.filter(
            ((GlobalChatMessage.sender_id == user1_id) & (GlobalChatMessage.recipient_id == user2_id)) |
            ((GlobalChatMessage.sender_id == user2_id) & (GlobalChatMessage.recipient_id == user1_id))
        ).order_by(GlobalChatMessage.sent_at).all()
        return jsonify([
            {
                'id': m.id,
                'sender_id': m.sender_id,
                'recipient_id': m.recipient_id,
                'content': m.content,
                'sent_at': m.sent_at.isoformat() if m.sent_at else None
            } for m in msgs
        ])
    # POST: send message
    data = request.get_json()
    content = data.get('content', '').strip()
    if not content:
        return jsonify({'error': 'Empty message'}), 400
    msg = GlobalChatMessage(
        id=str(uuid.uuid4()),
        sender_id=user1_id,
        recipient_id=user2_id,
        content=content
    )
    db.session.add(msg)
    db.session.commit()
    return jsonify({'msg': 'sent', 'id': msg.id})

@app.route('/global-chats', methods=['GET'])
@jwt_required()
def global_chats():
    current_user = get_jwt_identity()
    # Get all messages where user is sender or recipient
    msgs = GlobalChatMessage.query.filter(
        (GlobalChatMessage.sender_id == current_user) | (GlobalChatMessage.recipient_id == current_user)
    ).order_by(GlobalChatMessage.sent_at.desc()).all()
    # Map (other_user_id) -> latest message
    chat_map = {}
    for m in msgs:
        other_id = m.recipient_id if m.sender_id == current_user else m.sender_id
        if other_id not in chat_map:
            chat_map[other_id] = m
    # Get user info for all other users
    user_ids = list(chat_map.keys())
    users = User.query.filter(User.id.in_(user_ids)).all()
    user_dict = {u.id: {'id': u.id, 'name': u.name, 'email': u.email, 'avatar_url': u.avatar_url} for u in users}
    # Compose result
    result = []
    for other_id, msg in chat_map.items():
        result.append({
            'user': user_dict.get(other_id, {'id': other_id, 'name': 'Unknown', 'email': '', 'avatar_url': None}),
            'last_message': {
                'id': msg.id,
                'sender_id': msg.sender_id,
                'recipient_id': msg.recipient_id,
                'content': msg.content,
                'sent_at': msg.sent_at.isoformat() if msg.sent_at else None
            }
        })
    # Sort by latest message
    result.sort(key=lambda x: x['last_message']['sent_at'], reverse=True)
    return jsonify(result)

# -- Discovery & Suggestions --
@app.route('/discovery', methods=['GET'])
@jwt_required()
def discovery(): return jsonify([vars(i) for i in DiscoveryItem.query.all()])

@app.route('/users/<id>/suggestions', methods=['GET'])
@jwt_required()
def suggestions(id): return jsonify([vars(s) for s in CollaborationSuggestion.query.filter_by(for_user_id=id)])

# -- Grants & Applications --
@app.route('/grants', methods=['GET','POST'])
@jwt_required()
def grants_route():
    if request.method == 'GET':
        grants = Grant.query.all()
        return jsonify([{
            'id': g.id,
            'title': g.title,
            'description': g.description,
            'total_funding_usd': str(g.total_funding_usd) if g.total_funding_usd is not None else None,
            'application_questions': g.application_questions,
            'created_by_id': g.created_by_id,
            'created_at': g.created_at.isoformat() if g.created_at else None
        } for g in grants])
    elif request.method == 'POST':
        data = request.get_json()
        # Remove created_by_id from data if present to avoid duplicate kwarg
        data.pop('created_by_id', None)
        g = Grant(id=db.func.gen_random_uuid(), created_by_id=get_jwt_identity(), **data)
        db.session.add(g)
        db.session.commit()
        return jsonify({'id': g.id}), 201

@app.route('/grants/<id>', methods=['GET','POST'])
@jwt_required()
def get_grant(id):
    if request.method=='GET':
        g = Grant.query.get(id)
        if not g:
            return jsonify({'error': 'Grant not found'}), 404
        return jsonify({
            'id': g.id,
            'title': g.title,
            'description': g.description,
            'total_funding_usd': str(g.total_funding_usd) if g.total_funding_usd is not None else None,
            'application_questions': g.application_questions,
            'created_by_id': g.created_by_id,
            'created_at': g.created_at.isoformat() if g.created_at else None
        })
    data=request.get_json(); g=Grant(id=db.func.gen_random_uuid(), created_by_id=get_jwt_identity(), **data)
    db.session.add(g); db.session.commit(); return jsonify(id=g.id),201


@app.route('/grants/<id>/apply', methods=['POST'])
@jwt_required()
def apply_grant(id):
    data=request.get_json()
    ga=GrantApplication(id=db.func.gen_random_uuid(), grant_id=id, project_id=data['project_id'], answers=data['answers'])
    db.session.add(ga); db.session.commit(); return jsonify(id=ga.id),201

@app.route('/grants/<id>/applicants', methods=['GET'])
@jwt_required()
def applicants(id):
    return jsonify([a.to_dict() for a in GrantApplication.query.filter_by(grant_id=id)])

def is_valid_uuid(val):
    uuid_regex = re.compile(r'^[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}$')
    return bool(uuid_regex.match(str(val)))

@app.route('/grant-applications', methods=['POST'])
@jwt_required()
def create_grant_application():
    data = request.get_json()
    grant_id = data.get('grant_id')
    grant_name = data.get('grant')
    if not grant_id and grant_name:
        grant = Grant.query.filter_by(title=grant_name).first()
        if not grant:
            return jsonify({'error': 'Grant not found', 'debug': {'grant_name': grant_name}}), 404
        grant_id = grant.id
    if not grant_id:
        return jsonify({'error': 'Grant ID required', 'debug': {'grant_id': grant_id, 'grant_name': grant_name}}), 400
    project_id = data.get('project_id')
    if not project_id:
        return jsonify({'error': 'Project ID required', 'debug': {'project_id': project_id}}), 400
    # If project_id is not a valid UUID, only try to look up by title (case-insensitive)
    if not is_valid_uuid(project_id):
        user_id = get_jwt_identity()
        project = Project.query.filter(db.func.lower(Project.title) == project_id.lower(), Project.owner_id == user_id).first()
        if not project:
            return jsonify({'error': 'Project not found or invalid project_id', 'debug': {'project_id': project_id}}), 400
        project_id = project.id
    answers = data.get('answers')
    if not answers:
        answers = {
            'fit': data.get('fit'),
            'usage': data.get('usage'
        )}
    ga = GrantApplication(
        id=str(uuid.uuid4()),
        grant_id=grant_id,
        project_id=project_id,
        answers=answers
    )
    db.session.add(ga)
    db.session.commit()
    return jsonify({'id': ga.id}), 201

# -- Awards & Milestones --
@app.route('/applications/<aid>/award', methods=['POST'])
@jwt_required()
def award(aid):
    data=request.get_json(); aw=Award(id=db.func.gen_random_uuid(), application_id=aid, awarded_amount=data['awarded_amount'])
    db.session.add(aw); db.session.commit(); return jsonify(id=aw.id),201

@app.route('/awards/<id>/milestones', methods=['GET','POST'])
@jwt_required()
def milestones(id):
    if request.method=='GET': return jsonify([vars(m) for m in GrantMilestone.query.filter_by(award_id=id)])
    data=request.get_json(); m=GrantMilestone(id=db.func.gen_random_uuid(), award_id=id, name=data['name'], due_date=data.get('due_date'))
    db.session.add(m); db.session.commit(); return jsonify(id=m.id),201

# -- Projects CRUD --
@app.route('/projects', methods=['GET','POST'])
@jwt_required()
def projects():
    if request.method=='GET':
        return jsonify([{
            'id': p.id,
            'title': p.title,
            'owner_id': p.owner_id,
            'budget_requested': float(p.budget_requested) if p.budget_requested is not None else None,
            'reproducibility_score': float(p.reproducibility_score) if p.reproducibility_score is not None else None,
            'impact_score': float(p.impact_score) if p.impact_score is not None else None,
            'difficulty_score': float(p.difficulty_score) if p.difficulty_score is not None else None,
            'created_at': p.created_at.isoformat() if p.created_at else None,
            'updated_at': p.updated_at.isoformat() if p.updated_at else None,
        } for p in Project.query.all()])
    data = request.get_json()
    p = Project(
        id=db.func.gen_random_uuid(),
        owner_id=data['owner_id'],
        title=data['title'],
        budget_requested=None,  # AI will fill this later
        reproducibility_score=None,
        impact_score=None,
        difficulty_score=None,
    )
    db.session.add(p)
    # Also create a matching Experiment
    exp = Experiment(
        id=p.id,
        title=data['title'],
        description=data.get('description'),
        owner_id=data['owner_id'],
        visibility='public',
    )
    db.session.add(exp)
    db.session.flush()  # Ensure exp.id is available
    # Create initial protocol version for the experiment
    initial_version = ProtocolVersion(
        id=str(uuid.uuid4()),
        experiment_id=exp.id,
        version_label='v1.0',
        parent_version_id=None,
        metadata_={},
    )
    db.session.add(initial_version)
    db.session.commit()
    return jsonify(id=p.id), 201

@app.route('/projects/<id>', methods=['GET'])
@jwt_required()
def get_project(id):
    p = Project.query.get(id)
    if not p:
        abort(404, description='Project not found')
    return jsonify({
        'id': p.id,
        'title': p.title,
        'owner_id': p.owner_id,
        'budget_requested': str(p.budget_requested) if p.budget_requested else None,
        'reproducibility_score': p.reproducibility_score,
        'impact_score': p.impact_score,
        'difficulty_score': p.difficulty_score,
        'created_at': p.created_at.isoformat() if p.created_at else None,
        'updated_at': p.updated_at.isoformat() if p.updated_at else None,
    })

@app.route('/projects/<id>/paper', methods=['GET', 'PUT'])
@jwt_required()
def project_paper(id):
    p = Project.query.get(id)
    if not p:
        abort(404, description='Project not found')
    if request.method == 'GET':
        return jsonify({'content': p.paper_content or ''})
    data = request.get_json()
    p.paper_content = data.get('content', '')
    db.session.commit()
    return jsonify({'msg': 'updated'})

@app.route('/projects/<id>/grade', methods=['POST'])
@jwt_required()
def grade_experiment(id):
    project = Project.query.get(id)
    if not project:
        return jsonify({'error': 'Project not found'}), 404
    # Find experiment with same title and owner_id
    experiment = Experiment.query.filter_by(title=project.title, owner_id=project.owner_id).first()
    if not experiment:
        return jsonify({'error': 'Experiment not found for this project'}), 404
    # Get all protocol versions (use latest)
    versions = ProtocolVersion.query.filter_by(experiment_id=experiment.id).order_by(ProtocolVersion.created_at.desc()).all()
    protocol = versions[0] if versions else None
    # Get all steps for this protocol version
    steps = ExperimentStep.query.filter_by(protocol_version_id=protocol.id).order_by(ExperimentStep.order_index).all() if protocol else []
    step_texts = [f"Step {s.order_index+1}: {s.title}\n{s.content_markdown or ''}" for s in steps]
    # Get all attachments for all steps
    attachments = FileAttachment.query.filter(FileAttachment.experiment_step_id.in_([s.id for s in steps])).all() if steps else []
    attachment_texts = []
    for att in attachments:
        path = os.path.join(UPLOAD_FOLDER, att.storage_path)
        try:
            if att.filename.lower().endswith('.pdf'):
                with open(path, 'rb') as f:
                    reader = PdfReader(f)
                    text = "\n".join(page.extract_text() or '' for page in reader.pages)
                    attachment_texts.append(f"Attachment: {att.filename}\n{text}")
            elif att.filename.lower().endswith('.docx'):
                doc = Document(path)
                text = "\n".join([p.text for p in doc.paragraphs])
                attachment_texts.append(f"Attachment: {att.filename}\n{text}")
        except Exception as e:
            attachment_texts.append(f"Attachment: {att.filename} [Could not extract text: {e}]")
    # Get step map/flow
    step_map = protocol.step_map if protocol and protocol.step_map else {}
    # Compose prompt
    prompt = f"""
You are an expert scientific reviewer at a top‐10 research university with unlimited, world‐class core facilities and experienced staff. All experiments occur in labs equipped with the latest instrumentation and staffed by PhD‐level operators. Given the following experiment details, provide:

  • A reproducibility score (1.00–10.00, two nonzero decimal places)  
  • An impact score (1.00–10.00, two nonzero decimal places)  
  • A difficulty score (1.00–10.00, two nonzero decimal places)  
  • An estimated budget in USD (integer, no “$” sign, no commas)

Use these detailed scoring anchors:

  • Reproducibility (1.00 – 10.00):  
      – 1.00 = “Essentially impossible to replicate, even with full methods and top equipment.”  
      – 10.00 = “Trivial to reproduce in any similar state‐of‐the‐art lab with standard practice.”  

  • Impact (1.00 – 10.00):  
      – 1.00 = “Purely incremental or routine; no measurable advance in the field.”  
      – 10.00 = “A paradigm‐shifting result or technology that fundamentally changes how researchers approach this problem.”  

  • Difficulty (1.00 – 10.00):  
      – 1.00 = “Easily done in any core facility using off‐the‐shelf kits and minimal specialist training.”  
      – 3.00 = “Basic molecular biology or instrumentation tasks (e.g., routine cell culture, simple Western blot).”  
      – 5.00 = “Moderately challenging tasks requiring specialized know‐how (e.g., multistep purifications, advanced microscopy).”  
      – 7.00 = “Very challenging tasks even for top labs (e.g., developing a novel microfluidics device, custom animal model).”  
      – 10.00 = “Almost impossible without a dedicated, multi‐year team effort (e.g., inventing a new modality of single‐molecule detection).”  

      Estimate the realistic, minimal cost of executing this experiment without bloated institutional overhead. Focus only on core expenses: materials, equipment use, services, and personnel time (e.g., grad student, technician). Assume access to a well-resourced lab but estimate what must actually be spent to make the experiment work.

        Budget tiers (for guidance only):

        1–100 USD = “Basic proof-of-concept or classroom-level experiments (e.g., household materials, surveys, simple electronics, toy simulations).”

        100–2,000 USD = “Standard lab work or pilot studies (e.g., gel electrophoresis, Arduino prototypes, psych experiments with Amazon MTurk, basic behavioral assays).”

        2,000–10,000 USD = “Moderate complexity (e.g., small sequencing runs, prototype hardware builds, cell cultures, ML model training with paid compute).”

        10,000–100,000 USD = “Advanced or scaled setups (e.g., animal models, cleanroom fabrication, advanced neuroimaging, quantum optics).”

        >100,000 USD = “Extensive, multidisciplinary, or multi-phase programs (e.g., clinical trials, custom robotics platforms, satellite payloads, large HPC runs).”

        Avoid inflating for PI salaries, indirect costs, or facility markups unless absolutely central to execution.
      
Do NOT use round numbers or round decimals for the three scores (e.g., avoid “5.00” or “7.50”). All scores must have two nonzero decimal places (e.g., 4.23, 8.17). For budget, return an integer (e.g., 42500).

Example (for calibration only):

  Experiment: “High‐Throughput Single‐Cell RNA‐Seq of Drug‐Treated Cancer Cells”  
  → {{  
      "reproducibility": 8.42,  
      "impact": 7.15,  
      "difficulty": 6.38,  
      "estimated_budget": 78000  
    }}

---

Experiment Name: {experiment.title}

Experiment Description:  
{experiment.description or ''}

All Steps (in order):  
{chr(10).join(step_texts)}

Relevant Attachments (if any):  
{chr(10).join(attachment_texts)}

Step Map (JSON):  
{step_map}

Respond ONLY with a JSON object and keys exactly:  
```json
{{
  "reproducibility": float,
  "impact": float,
  "difficulty": float,
  "estimated_budget": int
}}
```
"""
    openai.api_key = "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"
    try:
        response = openai.chat.completions.create(
            model="o3-mini",
            messages=[{"role": "user", "content": prompt}],

        )
        import json
        content = response.choices[0].message.content
        import re
        match = re.search(r'{[\s\S]*}', content)
        if match:
            scores = json.loads(match.group(0))
            # Clamp and round
            for k in ["reproducibility","impact","difficulty"]:
                v = float(scores.get(k, 1.0))
                v = max(1.0, min(10.0, round(v,2)))
                scores[k] = v
            if "estimated_budget" in scores:
                try:
                    scores["estimated_budget"] = int(str(scores["estimated_budget"]).replace(",", "").replace("$", ""))
                except Exception:
                    scores["estimated_budget"] = None
            # Save to DB
            project.reproducibility_score = scores["reproducibility"]
            project.impact_score = scores["impact"]
            project.difficulty_score = scores["difficulty"]
            project.budget_requested = scores["estimated_budget"]
            db.session.commit()
            return jsonify(scores)
        else:
            return jsonify({"error": "No JSON in response", "raw": content}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/')
def index():
    return send_from_directory('.', 'anonymous-homepage.html')

@app.route('/login')
def login_page():
    return send_from_directory('.', 'login.html')

@app.route('/register')
def register_page():
    return send_from_directory('.', 'register.html')

@app.route('/<path:filename>')
def static_files(filename):
    return send_from_directory('.', filename)

@app.route('/protocol-versions/<id>/step-map', methods=['GET', 'PATCH'])
@jwt_required()
def protocol_version_step_map(id):
    pv = get_model_or_404(ProtocolVersion, id)
    if request.method == 'GET':
        return jsonify({'step_map': pv.step_map or {}})
    data = request.get_json()
    pv.step_map = data.get('step_map', {})
    db.session.commit()
    return jsonify({'msg': 'updated', 'step_map': pv.step_map})

@app.route('/projects/<id>/experiment', methods=['GET'])
@jwt_required()
def get_experiment_for_project(id):
    proj = Project.query.get(id)
    if not proj:
        abort(404, description='Project not found')
    # Find experiment with same title and owner_id
    exp = Experiment.query.filter_by(title=proj.title, owner_id=proj.owner_id).first()
    if not exp:
        abort(404, description='Experiment not found for this project')
    print(f"[DEBUG] Found experiment {exp.id} for project {id} at {datetime.datetime.now()}", flush=True)
    return jsonify({
        'id': exp.id,
        'title': exp.title,
        'description': exp.description,
        'owner_id': exp.owner_id,
        'visibility': exp.visibility,
        'created_at': exp.created_at.isoformat() if exp.created_at else None,
        'updated_at': exp.updated_at.isoformat() if exp.updated_at else None,
    })

@app.route('/users/<id>', methods=['GET'])
@jwt_required()
def get_user(id):
    user = User.query.filter_by(id=id).first()
    if not user:
        return jsonify({'error': 'User not found'}), 404
    return jsonify({
        'id': user.id,
        'email': user.email,
        'name': getattr(user, 'name', None) or user.email.split('@')[0]
    })

@app.route('/grant-applications/<id>', methods=['PATCH'])
@jwt_required()
def update_grant_application(id):
    app_obj = GrantApplication.query.get(id)
    if not app_obj:
        return jsonify({'error': 'Application not found'}), 404
    data = request.get_json()
    if 'status' in data:
        app_obj.status = data['status']
    db.session.commit()
    return jsonify({'msg': 'updated', 'id': app_obj.id, 'status': app_obj.status})

@app.route('/grant-applications', methods=['GET'])
@jwt_required()
def get_grant_applications():
    user_id = get_jwt_identity()
    user = User.query.get(user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
    # Funders see all applications for their grants, scientists see their own
    if user.role == 'funder':
        grants = Grant.query.filter_by(created_by_id=user_id).all()
        grant_ids = [g.id for g in grants]
        applications = GrantApplication.query.filter(GrantApplication.grant_id.in_(grant_ids)).all()
    elif user.role == 'scientist':
        projects = Project.query.filter_by(owner_id=user_id).all()
        project_ids = [p.id for p in projects]
        applications = GrantApplication.query.filter(GrantApplication.project_id.in_(project_ids)).all()
    else:
        applications = GrantApplication.query.all()
    # Preload grants and projects for lookup
    grant_map = {g.id: g for g in Grant.query.all()}
    project_map = {p.id: p for p in Project.query.all()}
    result = []
    for app in applications:
        grant = grant_map.get(app.grant_id)
        project = project_map.get(app.project_id)
        result.append({
            'id': app.id,
            'grant_id': app.grant_id,
            'project_id': app.project_id,
            'status': app.status,
            'answers': app.answers,
            'submitted_at': app.submitted_at.isoformat() if app.submitted_at else None,
            'updated_at': app.updated_at.isoformat() if app.updated_at else None,
            'application_questions': grant.application_questions if grant else [],
            'grant_title': grant.title if grant else None,
            'project_title': project.title if project else None,
            'applicant_id': project.owner_id if project else None
        })
    return jsonify(result)

@app.route('/grant-applications/<id>', methods=['GET'])
@jwt_required()
def get_grant_application(id):
    app_obj = GrantApplication.query.get(id)
    if not app_obj:
        return jsonify({'error': 'Not found'}), 404
    grant = Grant.query.get(app_obj.grant_id)
    project = Project.query.get(app_obj.project_id)
    return jsonify({
        'id': app_obj.id,
        'grant_id': app_obj.grant_id,
        'project_id': app_obj.project_id,
        'status': app_obj.status,
        'answers': app_obj.answers,
        'submitted_at': app_obj.submitted_at.isoformat() if app_obj.submitted_at else None,
        'updated_at': app_obj.updated_at.isoformat() if app_obj.updated_at else None,
        'application_questions': grant.application_questions if grant else [],
        'grant_title': grant.title if grant else None,
        'project_title': project.title if project else None,
        'applicant_id': project.owner_id if project else None
    })

@app.route('/steps/<id>/grade', methods=['POST'])
@jwt_required()
def grade_step(id):
    step = get_model_or_404(ExperimentStep, id)
    # Get protocol version and experiment
    protocol = step.protocol_version
    experiment = protocol.experiment if protocol else None
    # Get all previous steps (including this one), ordered
    steps = ExperimentStep.query.filter_by(protocol_version_id=step.protocol_version_id).order_by(ExperimentStep.order_index).all()
    step_texts = [f"Step {s.order_index+1}: {s.title}\n{s.content_markdown or ''}" for s in steps if s.order_index <= step.order_index]
    # Get experiment description
    experiment_desc = experiment.description if experiment else ''
    # Get all attachments for this step
    attachments = FileAttachment.query.filter_by(experiment_step_id=id).all()
    attachment_texts = []
    for att in attachments:
        path = os.path.join(UPLOAD_FOLDER, att.storage_path)
        try:
            if att.filename.lower().endswith('.pdf'):
                with open(path, 'rb') as f:
                    reader = PdfReader(f)
                    text = "\n".join(page.extract_text() or '' for page in reader.pages)
                    attachment_texts.append(f"Attachment: {att.filename}\n{text}")
            elif att.filename.lower().endswith('.docx'):
                doc = Document(path)
                text = "\n".join([p.text for p in doc.paragraphs])
                attachment_texts.append(f"Attachment: {att.filename}\n{text}")
        except Exception as e:
            attachment_texts.append(f"Attachment: {att.filename} [Could not extract text: {e}]")

    prompt = f"""
    You are an expert scientific reviewer at a top‐10 research university with unlimited, world‐class core facilities and experienced staff. All steps occur in labs equipped with the latest instrumentation and staffed by PhD‐level operators. Use these anchors:

    • Reproducibility (1.00 – 10.00):  
        – 1.00 = “Essentially impossible to replicate even with full methods and top equipment.”  
        – 10.00 = “Trivial to reproduce in any similar state of the art lab with standard practice.”  

    • Impact (1.00 – 10.00):  
        – 1.00 = “Purely incremental or routine; no measurable advance in the field.”  
        – 10.00 = “A paradigm‐shifting result or technology that fundamentally changes how researchers approach this problem.”  

    • Difficulty (1.00 – 10.00):  
        – 1.00 = “Easily done in any core facility using off‐the‐shelf kits and standard training.”  
        – 3.00 = “Basic molecular biology or instrumentation tasks (e.g., routine cloning, Western blot).”  
        – 5.00 = “Moderately challenging tasks requiring specialized know‐how (e.g., multistep purifications, advanced imaging).”  
        – 7.00 = “Very challenging tasks even for top labs (e.g., building a new custom microfluidics device from scratch).”  
        – 10.00 = “Almost impossible without a dedicated team over multiple years (e.g., inventing a brand-new modality of single-molecule detection).”  

    Do NOT use round numbers or round decimals (e.g., avoid “5.00,” “7.50”). All scores must have two nonzero decimal places (e.g., 4.23, 8.17).

    Example (for calibration only):  
    Step: “Perform PCR amplification of GFP using published primers.”  
    → {{ "reproducibility": 9.25, "impact": 1.47, "difficulty": 2.11 }}

    ---

    Experiment Description:  
    {experiment_desc}

    Previous and Current Steps:  
    {chr(10).join(step_texts)}

    Relevant Attachments (if any):  
    {chr(10).join(attachment_texts)}

    Respond ONLY with a JSON object and keys exactly:  
    ```json
    {{
    "reproducibility": float,
    "impact": float,
    "difficulty": float
    }}
    ```
    """


    # Call OpenAI
    openai.api_key = "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"
    try:
        response = openai.chat.completions.create(
            model="o3-mini",
            messages=[{"role": "user", "content": prompt}]
        )
        import json
        # Try to extract JSON from response
        content = response.choices[0].message.content
        # Find first { ... } block
        import re
        match = re.search(r'{[\s\S]*}', content)
        if match:
            scores = json.loads(match.group(0))
            # Clamp and round
            for k in ["reproducibility","impact","difficulty"]:
                v = float(scores.get(k, 1.0))
                v = max(1.0, min(10.0, round(v,2)))
                scores[k] = v
            # Save to DB
            step.reproducibility_score = scores["reproducibility"]
            step.impact_score = scores["impact"]
            step.difficulty_score = scores["difficulty"]
            db.session.commit()
            return jsonify(scores)
        else:
            return jsonify({"error": "No JSON in response", "raw": content}), 500
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# --- Lab Notebook API ---
@app.route('/projects/<project_id>/notebook', methods=['GET', 'POST', 'PATCH'])
@jwt_required()
def notebook_entries(project_id):
    user_id = get_jwt_identity()
    if request.method == 'GET':
        entries = NotebookEntry.query.filter_by(project_id=project_id).order_by(NotebookEntry.timestamp.desc()).all()
        return jsonify([e.to_dict() for e in entries])
    data = request.get_json()
    if request.method == 'POST':
        # New entry
        import uuid
        entry = NotebookEntry(
            id=str(uuid.uuid4()),
            project_id=project_id,
            user_id=user_id,
            user_name=data.get('user_name'),
            timestamp=datetime.datetime.now(),
            device=data.get('device'),
            location=data.get('location'),
            session_id=data.get('session_id'),
            experiment_id=data.get('experiment_id'),
            version=data.get('version'),
            visibility=data.get('visibility'),
            content=data.get('content'),
            structured=data.get('structured'),
            diffs=[]
        )
        db.session.add(entry)
        db.session.commit()
        return jsonify(entry.to_dict()), 201
    if request.method == 'PATCH':
        # Edit entry (with version diff)
        entry = NotebookEntry.query.get(data.get('id'))
        if not entry: return jsonify({'error': 'Not found'}), 404
        old_content = entry.content or ''
        new_content = data.get('content', old_content)
        if new_content != old_content:
            diff = '\n'.join(difflib.unified_diff(old_content.splitlines(), new_content.splitlines(), lineterm=''))
            entry.diffs = (entry.diffs or []) + [{
                'timestamp': datetime.datetime.now().isoformat(),
                'diff': diff
            }]
        entry.content = new_content
        entry.structured = data.get('structured', entry.structured)
        entry.device = data.get('device', entry.device)
        entry.location = data.get('location', entry.location)
        entry.visibility = data.get('visibility', entry.visibility)
        db.session.commit()
        return jsonify(entry.to_dict())

@app.route('/notebook-entries/<entry_id>/attachments', methods=['POST'])
@jwt_required()
def upload_notebook_attachment(entry_id):
    entry = NotebookEntry.query.get(entry_id)
    if not entry: return jsonify({'error': 'Not found'}), 404
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    import uuid
    ext = os.path.splitext(file.filename)[1]
    unique_id = str(uuid.uuid4())
    storage_filename = f"{unique_id}{ext}"
    storage_path = os.path.join('uploads', storage_filename)
    file.save(storage_path)
    att = NotebookAttachment(
        id=unique_id,
        entry_id=entry_id,
        filename=file.filename,
        storage_path=storage_filename
    )
    db.session.add(att)
    db.session.commit()
    return jsonify(att.to_dict()), 201

@app.route('/notebook-attachments/<attachment_id>/download', methods=['GET'])
def download_notebook_attachment(attachment_id):
    att = NotebookAttachment.query.get(attachment_id)
    if not att: return '', 404
    return send_from_directory('uploads', att.storage_path, as_attachment=True, download_name=att.filename)

@app.route('/projects/<project_id>/notebook/summary', methods=['GET'])
@jwt_required()
def notebook_summary(project_id):
    style = request.args.get('style', 'verbose')
    entries = NotebookEntry.query.filter_by(project_id=project_id).order_by(NotebookEntry.timestamp.desc()).all()
    text = '\n'.join([e.content or '' for e in entries])
    # Call LLM for summary
    import openai
    openai.api_key = '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'
    prompt = f"Summarize the following lab notebook entries in {style} style:\n{text}"
    try:
        resp = openai.chat.completions.create(
            model='gpt-3.5-turbo',
            messages=[{'role':'user','content':prompt}]
        )
        summary = resp.choices[0].message.content
    except Exception as e:
        summary = f"[AI summary unavailable: {e}]"
    return jsonify({'summary': summary})

# --- Journal Matching API ---
@app.route('/journal-match', methods=['POST'])
@jwt_required()
def journal_match():
    data = request.get_json()
    content = data.get('content', '')
    # Compose prompt for GPT
    prompt = f"""
Given the following scientific manuscript, suggest 3-6 journals that would be a good fit for submission. For each journal, provide:
- Name
- Short description (1-2 lines)
- URL (homepage or submission page)

Manuscript:
{content}

Respond as a JSON array of objects with keys: name, description, url.
"""
    import openai
    openai.api_key = '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'
    try:
        resp = openai.chat.completions.create(
            model='gpt-3.5-turbo',
            messages=[{'role':'user','content':prompt}]
        )
        import json, re
        content = resp.choices[0].message.content
        match = re.search(r'\[.*\]', content, re.DOTALL)
        journals = json.loads(match.group(0)) if match else []
    except Exception as e:
        journals = []
    return jsonify({'journals': journals})



# --- Copilot Chat API ---
@app.route('/copilot-chat', methods=['POST'])
@jwt_required()
def copilot_chat():
    data = request.get_json()
    message = data.get('message', '')
    project_id = request.args.get('project_id') or request.headers.get('X-Project-Id')
    # Gather all context: steps, results, ratings, notebook posts
    steps = []
    results = []
    ratings = []
    notebook = []
    if project_id:
        from sqlalchemy import or_
        # Steps
        steps = [s.to_dict() for s in ExperimentStep.query.join(ProtocolVersion, ExperimentStep.protocol_version_id==ProtocolVersion.id)
                 .filter(ProtocolVersion.experiment_id==project_id).all()]
        # Notebook
        notebook = [n.to_dict() for n in NotebookEntry.query.filter_by(project_id=project_id).all()]
    # Compose context
    context = f"Steps:\n{steps}\n\nNotebook Entries:\n{notebook}\n\n" \
              f"User Message:\n{message}\n"
    prompt = f"""
You are Atlantis Copilot, an expert scientific assistant. Use the following context from the experiment, steps, results, ratings, and lab notebook to answer the user's question. Be concise, helpful, and cite relevant step numbers or notebook entries if possible.\n\n{context}\n\nReply as a helpful assistant.\n"""
    import openai
    openai.api_key = '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'
    try:
        resp = openai.chat.completions.create(
            model='gpt-3.5-turbo',
            messages=[{'role':'user','content':prompt}]
        )
        reply = resp.choices[0].message.content
    except Exception as e:
        reply = f"[Copilot unavailable: {e}]"
    return jsonify({'reply': reply})

# --- Lab Model (assumed from lab_schema_migration.sql) ---
# Example fields: id, name, description, pi_name, institution, created_at, updated_at
# If your Lab model is in codex_api.py, import it; else, define here for reference.
class Lab(db.Model):
    __tablename__ = 'labs'
    id = db.Column(db.String, primary_key=True, default=lambda: str(uuid.uuid4()))
    name = db.Column(db.String(255), unique=True, nullable=False)
    description = db.Column(db.Text)
    affiliation = db.Column(db.String(255))
    created_by = db.Column(db.String, db.ForeignKey('users.id'), nullable=False)
    created_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())
    # Optionally, add relationships for members, etc.

# --- CRUD Endpoints ---

@app.route('/labs-list', methods=['GET'])
def list_labs():
    print("[DEBUG] /labs-list GET called", flush=True)
    try:
        labs = Lab.query.order_by(Lab.created_at.desc()).all()
        print(f"[DEBUG] /labs-list GET: found {len(labs)} labs", flush=True)
        return jsonify([{
            'id': l.id,
            'name': l.name,
            'description': l.description,
            'affiliation': l.affiliation,
            'created_by': l.created_by,
            'created_at': l.created_at.isoformat() if l.created_at else None
        } for l in labs])
    except Exception as e:
        import traceback
        print(f"[ERROR] /labs-list GET: {e}\n" + traceback.format_exc(), flush=True)
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500

@app.route('/lab-create', methods=['POST'])
@jwt_required()
def create_lab():
    print("[DEBUG] /lab-create POST called", flush=True)
    data = request.get_json()
    user_id = get_jwt_identity()
    lab = Lab(
        id=str(uuid.uuid4()),
        name=data['name'],
        description=data.get('description'),
        affiliation=data.get('affiliation'),
        created_by=user_id,
        created_at=datetime.datetime.now()
    )
    db.session.add(lab)
    db.session.commit()
    return jsonify({'id': lab.id}), 201

@app.route('/lab/<id>', methods=['GET'])
def get_lab(id):
    print(f"[DEBUG] /lab/{{id}} GET called", flush=True)
    lab = get_model_or_404(Lab, id)
    # Query members
    members = LabMember.query.filter_by(lab_id=lab.id).all()
    member_list = [{
        'user_id': m.user_id,
        'role': m.role
    } for m in members]
    # Find experiments for all members
    member_user_ids = [m.user_id for m in members]
    experiments = []
    if member_user_ids:
        exps = Experiment.query.filter(Experiment.owner_id.in_(member_user_ids)).all()
        experiments = [{
            'id': e.id,
            'title': e.title,
            'owner_id': e.owner_id
        } for e in exps]
    return jsonify({
        'id': lab.id,
        'name': lab.name,
        'description': lab.description,
        'affiliation': lab.affiliation,
        'created_by': lab.created_by,
        'created_at': lab.created_at.isoformat() if lab.created_at else None,
        'members': member_list,
        'experiments': experiments
    })

@app.route('/lab-update/<id>', methods=['PATCH'])
def update_lab(id):
    print(f"[DEBUG] /lab-update/{{id}} PATCH called", flush=True)
    lab = get_model_or_404(Lab, id)
    data = request.get_json()
    for k in ['name', 'description', 'affiliation']:
        if k in data:
            setattr(lab, k, data[k])
    db.session.commit()
    return jsonify({'msg': 'updated'})

@app.route('/lab-delete/<id>', methods=['DELETE'])
def delete_lab(id):
    print(f"[DEBUG] /lab-delete/{{id}} DELETE called", flush=True)
    lab = get_model_or_404(Lab, id)
    db.session.delete(lab)
    db.session.commit()
    return '', 204

# --- Lab Members Endpoints ---

class LabMember(db.Model):
    __tablename__ = 'lab_members'
    lab_id = db.Column(db.String, db.ForeignKey('labs.id'), primary_key=True)
    user_id = db.Column(db.String, db.ForeignKey('users.id'), primary_key=True)
    role = db.Column(db.String(50), default='member')
    joined_at = db.Column(db.DateTime(timezone=True), server_default=db.func.now())

@app.route('/lab/<lab_id>/members', methods=['POST'])
@jwt_required()
def add_lab_member(lab_id):
    print(f"[DEBUG] /lab/{{lab_id}}/members POST called", flush=True)
    data = request.get_json()
    user_id = data.get('user_id')
    role = data.get('role', 'member')
    if not user_id:
        return jsonify({'error': 'user_id required'}), 400
    # Check if already a member
    existing = LabMember.query.filter_by(lab_id=lab_id, user_id=user_id).first()
    if existing:
        return jsonify({'error': 'User already a member'}), 400
    member = LabMember(lab_id=lab_id, user_id=user_id, role=role)
    db.session.add(member)
    db.session.commit()
    return jsonify({'msg': 'added', 'user_id': user_id, 'role': role}), 201

@app.route('/projects/<project_id>/steps', methods=['GET'])
@jwt_required()
def project_steps(project_id):
    # Find all protocol versions for this project
    project_id = proj_to_experiment(project_id)
    protocol_versions = [ProtocolVersion.query.filter_by(experiment_id=project_id).order_by(ProtocolVersion.created_at.desc()).first()]
    steps = []
    for pv in protocol_versions:
        pv_steps = ExperimentStep.query.filter_by(protocol_version_id=pv.id).all()
        steps.extend([s.to_dict() for s in pv_steps])
    return jsonify(steps)

@app.route('/projects/<project_id>/fork', methods=['POST'])
@jwt_required()
def fork_project(project_id):
    user_id = get_jwt_identity()
    orig_project = Project.query.get(project_id)
    if not orig_project:
        return jsonify({'error': 'Project not found'}), 404
    if orig_project.owner_id == user_id:
        return jsonify({'error': 'You already own this project.'}), 400
    # Find matching experiment
    orig_experiment = Experiment.query.filter_by(title=orig_project.title, owner_id=orig_project.owner_id).first()
    if not orig_experiment:
        return jsonify({'error': 'Experiment not found for this project'}), 404
    # Create new project
    new_project_id = uuid.uuid4()
    new_project = Project(
        id=new_project_id,
        owner_id=user_id,
        title=orig_project.title,
        budget_requested=orig_project.budget_requested,
        reproducibility_score=orig_project.reproducibility_score,
        impact_score=orig_project.impact_score,
        difficulty_score=orig_project.difficulty_score,
        paper_content=getattr(orig_project, 'paper_content', None)
    )

    db.session.add(new_project)
    # Create new experiment
    new_experiment = Experiment(
        id=new_project_id,
        title=orig_experiment.title,
        description=orig_experiment.description,
        owner_id=user_id,
        visibility=orig_experiment.visibility,
    )
    db.session.add(new_experiment)
    db.session.flush()
    # Find latest protocol version
    orig_versions = ProtocolVersion.query.filter_by(experiment_id=orig_experiment.id).order_by(ProtocolVersion.created_at.desc()).all()
    if not orig_versions:
        return jsonify({'error': 'No protocol version to fork'}), 400
    orig_version = orig_versions[0]
    # Create new protocol version for fork
    new_version = ProtocolVersion(
        id=uuid.uuid4(),
        experiment_id=new_experiment.id,
        version_label='v1.0 (forked)',
        parent_version_id=orig_version.id,
        metadata_=orig_version.metadata_,
        step_map=orig_version.step_map
    )
    db.session.add(new_version)
    db.session.flush()
    # Copy steps
    orig_steps = ExperimentStep.query.filter_by(protocol_version_id=orig_version.id).all()
    for s in orig_steps:
        new_step = ExperimentStep(
            id=uuid.uuid4(),
            protocol_version_id=new_version.id,
            title=s.title,
            due_date=s.due_date,
            done=False,
            content_markdown=s.content_markdown,
            results_markdown=getattr(s, 'results_markdown', None),
            estimated_time_minutes=s.estimated_time_minutes,
            assigned_to_id=None,
            reproducibility_score=s.reproducibility_score,
            impact_score=s.impact_score,
            difficulty_score=s.difficulty_score,
            order_index=s.order_index
        )
        db.session.add(new_step)
    db.session.commit()
    return jsonify({'id': new_project_id}), 201

if __name__ == '__main__':
    # Create upload folder if it doesn't exist
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    # Run the Flask app
    app.run(debug=True, host='0.0.0.0', port=5000)